from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass
class ParsedDocument:
    raw_text: str
    ocr_text: str | None
    page_texts: list[str]


class DocumentParser:
    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix in {".png", ".jpg", ".jpeg", ".bmp"}:
            return self._parse_image(path)
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(raw_text=text, ocr_text=None, page_texts=[text])

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        try:
            import fitz  # type: ignore

            doc = fitz.open(path)
            pages = [page.get_text("text") for page in doc]
            return ParsedDocument(raw_text="\n".join(pages), ocr_text=None, page_texts=pages)
        except Exception:
            raw = path.read_bytes().decode("utf-8", errors="ignore")
            return ParsedDocument(raw_text=raw, ocr_text=None, page_texts=[raw])

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document  # type: ignore

            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ParsedDocument(raw_text=text, ocr_text=None, page_texts=[text])
        except Exception:
            raw = path.read_bytes().decode("utf-8", errors="ignore")
            return ParsedDocument(raw_text=raw, ocr_text=None, page_texts=[raw])

    def _parse_image(self, path: Path) -> ParsedDocument:
        try:
            import pytesseract  # type: ignore
            from PIL import Image

            text = pytesseract.image_to_string(Image.open(path), lang="chi_sim+eng")
            return ParsedDocument(raw_text=text, ocr_text=text, page_texts=[text])
        except Exception:
            return ParsedDocument(raw_text="", ocr_text="", page_texts=[""])

